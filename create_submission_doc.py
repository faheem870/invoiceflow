from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== TITLE =====
title = doc.add_heading('InvoiceFlow', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0xF0, 0xB9, 0x0B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('BNB Smart Builders Challenge — Hackathon Submission')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph()

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tagline.add_run('Turn real-world invoices into programmable on-chain receivables on BNB Chain')
run.font.size = Pt(12)
run.italic = True

doc.add_paragraph()

# ===== PROJECT OVERVIEW =====
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'InvoiceFlow is a non-custodial, open-source platform that tokenizes real-world invoices '
    'as NFTs on BNB Chain. It enables freelancers and SMEs to get paid in stablecoins via escrow, '
    'sell invoices early (factoring) to investors for instant cashflow, and leverages AI for risk '
    'scoring and fraud detection. It also includes a DeSci research pool for anonymized payment data.'
)

# ===== PROBLEM =====
doc.add_heading('2. Problem Statement', level=1)
doc.add_paragraph(
    'Freelancers and SMEs wait 30–90 days to get paid. Meanwhile, they cannot invest, grow, or '
    'cover operating costs. Traditional factoring is expensive (2-5% fees), opaque, and inaccessible '
    'to small businesses. Cross-border payments add further complexity and delays.'
)

# ===== SOLUTION =====
doc.add_heading('3. Solution', level=1)
bullets = [
    'Instant stablecoin payments via non-custodial escrow smart contracts',
    'Invoice factoring marketplace — sell receivables to investors at a discount for instant cash',
    'AI-powered risk scoring with automated discount suggestions and fraud detection',
    'DeSci research pool — opt-in anonymized payment data for academic research',
    'Full invoice lifecycle management: Draft → Approval → Payment/Listing → Settlement',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ===== THEME MAPPING =====
doc.add_heading('4. BNB Chain Theme Mapping', level=1)
doc.add_paragraph('InvoiceFlow covers ALL 7 BNB Chain ecosystem themes:')

table = doc.add_table(rows=8, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = table.rows[0].cells
headers[0].text = 'Theme'
headers[1].text = 'Implementation'
for cell in headers:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

data = [
    ('Payments', 'Stablecoin invoice settlement via escrow, payment links, cross-border friendly'),
    ('RWA', 'Invoices tokenized as ERC-721 NFTs — real-world receivables on-chain'),
    ('Trading', 'Invoice marketplace — buy/sell invoice NFTs with transparent pricing'),
    ('DeFi', 'Factoring = DeFi yield — investors buy invoices at discount, earn spread on settlement'),
    ('AI', 'AI agent for risk scoring, discount suggestions, fraud detection'),
    ('Wallets', 'MetaMask/WalletConnect integration, one-click stablecoin payments'),
    ('DeSci', 'Research pool with opt-in anonymized data, grant allocation, access NFTs'),
]
for i, (theme, impl) in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = theme
    row[1].text = impl

# ===== TECH STACK =====
doc.add_heading('5. Tech Stack', level=1)
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h = table2.rows[0].cells
h[0].text = 'Layer'
h[1].text = 'Technology'
for cell in h:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

stack = [
    ('Smart Contracts', 'Solidity 0.8.24, Hardhat, OpenZeppelin v5'),
    ('Frontend', 'Vue 3, Vite, TailwindCSS, ethers.js v6'),
    ('Backend', 'Express, Prisma, PostgreSQL'),
    ('AI', 'Risk scoring API (rule-based MVP, ML-ready)'),
    ('Chain', 'BNB Smart Chain (BSC Testnet)'),
]
for i, (layer, tech) in enumerate(stack):
    row = table2.rows[i + 1].cells
    row[0].text = layer
    row[1].text = tech

# ===== SMART CONTRACTS (LIVE PROOF) =====
doc.add_heading('6. Smart Contracts — Deployed on BSC Testnet ✅', level=1)
doc.add_paragraph('All 5 contracts are live and verified on BSC Testnet (Chain ID: 97):')

table3 = doc.add_table(rows=6, cols=3)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
h3 = table3.rows[0].cells
h3[0].text = 'Contract'
h3[1].text = 'Address'
h3[2].text = 'BscScan Link'
for cell in h3:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

contracts = [
    ('InvoiceNFT', '0x5e43b9C7910e938A1f3104E1552f002f22EDbf22'),
    ('InvoiceEscrow', '0x28E9ae4cBE146e90E99E748b6bb47234CFB383Dc'),
    ('InvoiceMarketplace', '0x2Cfa30007942020A0730bcb514f72a9Acd3387c4'),
    ('ResearchPool', '0x2dEc98F4fF7ADd31A67b73eC7ACa5bd97ee8DB98'),
    ('MockStablecoin', '0x1eA608aC5AF0130Cbb3dcD705797EAb56993E8Ca'),
]
for i, (name, addr) in enumerate(contracts):
    row = table3.rows[i + 1].cells
    row[0].text = name
    row[1].text = addr
    row[2].text = f'https://testnet.bscscan.com/address/{addr}'

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Deployer Address: ')
run.font.bold = True
p.add_run('0x011f48b8650784EEEe13DA4E95b98F5e03f9114e')

p2 = doc.add_paragraph()
run2 = p2.add_run('Deployer BscScan: ')
run2.font.bold = True
p2.add_run('https://testnet.bscscan.com/address/0x011f48b8650784EEEe13DA4E95b98F5e03f9114e')

# ===== ARCHITECTURE =====
doc.add_heading('7. Architecture', level=1)
doc.add_paragraph(
    'InvoiceFlow uses a modular architecture with 4 core smart contracts that interact '
    'with each other through authorized module patterns:'
)

arch_items = [
    'InvoiceNFT (ERC-721) — Core invoice token with full lifecycle status management (Draft → Awaiting Approval → Approved → Listed → Sold → Paid). Supports dispute resolution via arbitrator role.',
    'InvoiceEscrow — Handles stablecoin payments from payers to current invoice owners. Deducts protocol fees (0.3%) and optional research pool fees. Uses ReentrancyGuard for security.',
    'InvoiceMarketplace — DeFi factoring marketplace where sellers list approved invoices at a discount. Investors buy invoice NFTs and earn yield when the payer settles. Marketplace fee: 0.5%.',
    'ResearchPool (DeSci) — Accepts stablecoin donations, allocates grants to researchers via admin/DAO, and mints Access NFTs for data marketplace licenses.',
    'MockStablecoin (ERC-20) — Test USDT token for demo and testing purposes.',
]
for item in arch_items:
    doc.add_paragraph(item, style='List Bullet')

# ===== INVOICE LIFECYCLE =====
doc.add_heading('8. Invoice Lifecycle', level=1)
doc.add_paragraph('DRAFT → AWAITING_APPROVAL → APPROVED → LISTED → SOLD → PAID')
doc.add_paragraph('                                    ↓                              ↑')
doc.add_paragraph('                              (direct payment) ─────────────────────')
doc.add_paragraph('                    DISPUTED → APPROVED / CANCELLED')
doc.add_paragraph()
lifecycle = [
    'Seller creates invoice and mints as DRAFT NFT',
    'Seller requests payer approval → status becomes AWAITING_APPROVAL',
    'Payer reviews and approves (or disputes) → APPROVED',
    'Option A: Payer pays directly via escrow → PAID',
    'Option B: Seller lists on marketplace → LISTED → Investor buys → SOLD → Payer pays investor via escrow → PAID',
    'Disputes resolved by arbitrator (admin multisig for MVP)',
]
for item in lifecycle:
    doc.add_paragraph(item, style='List Number')

# ===== TESTS =====
doc.add_heading('9. Test Results — 18/18 Passing ✅', level=1)
doc.add_paragraph('Comprehensive test suite covering all contracts and end-to-end flows:')

table4 = doc.add_table(rows=6, cols=2)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
h4 = table4.rows[0].cells
h4[0].text = 'Test Suite'
h4[1].text = 'Tests'
for cell in h4:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

tests = [
    ('InvoiceNFT', '7 — mint, approve, dispute, resolve, cancel, edge cases'),
    ('InvoiceEscrow', '3 — pay approved invoice, reject unauthorized, reject non-payer'),
    ('InvoiceMarketplace', '3 — list+buy, cancel listing, reject non-approved'),
    ('ResearchPool (DeSci)', '3 — donate+grants, access NFTs, double-execute prevention'),
    ('End-to-End Flows', '2 — full payment lifecycle, full factoring lifecycle'),
]
for i, (suite, desc) in enumerate(tests):
    row = table4.rows[i + 1].cells
    row[0].text = suite
    row[1].text = desc

# ===== KEY FEATURES =====
doc.add_heading('10. Key Features', level=1)
features = [
    'Non-custodial — users maintain full control of their assets',
    'Open-source — MIT licensed, free for anyone to use and build upon',
    'Invoice NFTs — real-world receivables tokenized as ERC-721 tokens',
    'Stablecoin escrow — secure payment routing with on-chain receipts',
    'DeFi factoring — investors earn yield by purchasing discounted invoices',
    'Protocol fees — sustainable 0.3% settlement fee + 0.5% marketplace fee',
    'Dispute resolution — arbitrator role for handling payment disputes',
    'DeSci integration — research pool with donations, grants, and access NFTs',
    'AI risk scoring — automated discount suggestions based on payer history',
    'Multi-role support — Seller, Payer, Investor, Arbitrator, Researcher',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# ===== LINKS =====
doc.add_heading('11. Project Links', level=1)
table5 = doc.add_table(rows=5, cols=2)
table5.style = 'Light Grid Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
links = [
    ('GitHub Repository', 'https://github.com/faheem870/invoiceflow'),
    ('BSC Testnet Explorer', 'https://testnet.bscscan.com/address/0x011f48b8650784EEEe13DA4E95b98F5e03f9114e'),
    ('InvoiceNFT Contract', 'https://testnet.bscscan.com/address/0x5e43b9C7910e938A1f3104E1552f002f22EDbf22'),
    ('InvoiceEscrow Contract', 'https://testnet.bscscan.com/address/0x28E9ae4cBE146e90E99E748b6bb47234CFB383Dc'),
    ('InvoiceMarketplace Contract', 'https://testnet.bscscan.com/address/0x2Cfa30007942020A0730bcb514f72a9Acd3387c4'),
]
for i, (label, url) in enumerate(links):
    row = table5.rows[i].cells
    row[0].text = label
    row[1].text = url

# ===== ACCEPTANCE CRITERIA =====
doc.add_heading('12. Hackathon Acceptance Criteria', level=1)
criteria = [
    ('✅ Deployed contracts on BSC Testnet', 'All 5 contracts deployed with verified addresses'),
    ('✅ 2+ successful transactions', 'Mint NFT, Approve, List, Buy, Pay — all working on-chain'),
    ('✅ Explorer links in README', 'All contract addresses linked to testnet.bscscan.com'),
    ('✅ Open-source repo', 'MIT License, full source code on GitHub'),
    ('✅ Covers all 7 themes', 'Payments, RWA, DeFi, Trading, AI, Wallets, DeSci'),
    ('✅ 18/18 tests passing', 'Comprehensive smart contract test suite'),
]
for check, desc in criteria:
    p = doc.add_paragraph()
    run = p.add_run(check + ' — ')
    run.font.bold = True
    p.add_run(desc)

# ===== TEAM =====
doc.add_heading('13. Team', level=1)
p = doc.add_paragraph()
run = p.add_run('Faheem Khan')
run.font.bold = True
run.font.size = Pt(13)
doc.add_paragraph('AI Systems Architect & Full-Stack Developer')
team_details = [
    '15+ years in digital transformation',
    '1,000+ workflows deployed across enterprise clients',
    '50,000+ manual hours eliminated',
    'Expertise: MCP, CRM/ERP Integration, AI Agents, Voice AI, Web3',
    'Top Rated Plus on Upwork (Top 3%)',
    'LinkedIn: linkedin.com/in/faheem87',
    'Website: faheemkhan.co',
]
for d in team_details:
    doc.add_paragraph(d, style='List Bullet')

# ===== ROADMAP =====
doc.add_heading('14. Post-Hackathon Roadmap', level=1)
roadmap = [
    ('v1.1', 'Bulk/bundle invoice trading, partial payments, auto reminders'),
    ('v2.0', 'Liquidity pool + LP shares, account abstraction, gas sponsorship'),
    ('v3.0', 'On/off ramp partners, compliance & enterprise integrations'),
]
for ver, desc in roadmap:
    p = doc.add_paragraph()
    run = p.add_run(f'{ver}: ')
    run.font.bold = True
    p.add_run(desc)

# ===== FOOTER =====
doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Built with 💛 for BNB Smart Builders Challenge')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xF0, 0xB9, 0x0B)

# Save
output_path = '/Users/faheem/Desktop/Haya/invoiceflow/InvoiceFlow_Submission.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
